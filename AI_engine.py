import os
import serpapi
from docx import Document
from docx.shared import Inches, Pt, RGBColor
import httpx
from bs4 import BeautifulSoup
import json
import re
import matplotlib
matplotlib.use('Agg') 
import matplotlib.pyplot as plt
import pandas as pd
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image as RLImage, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet
import fitz 

from report_formats import get_template_instructions

# --- 2-LAYER MODEL CONFIGURATION ---
SMART_MODEL = "x-ai/grok-4.1-fast:free"
BACKUP_MODEL = "meta-llama/llama-3.3-70b-instruct:free"

SEARCH_RESULTS_COUNT = 10
MAX_RESULTS_TO_SCRAPE = 3
WORDS_PER_PAGE = 400

# --- HELPER FUNCTIONS ---
def clean_ai_output(text: str) -> str:
    if not text: return ""
    text = re.sub(r'<think>.*?</think>', '', text, flags=re.DOTALL)
    text = re.sub(r'^```\w*\s*', '', text, flags=re.MULTILINE)
    text = re.sub(r'^```\s*$', '', text, flags=re.MULTILINE)
    return text.strip()

def clean_section_output(text: str, section_title: str) -> str:
    if not text: return ""
    text = clean_ai_output(text)
    lines = text.split('\n')
    while lines and not lines[0].strip(): lines.pop(0)
    if not lines: return ""
    first_line = lines[0].strip().lower()
    clean_title = section_title.lower().replace('#', '').strip()
    clean_first = first_line.replace('#', '').strip()
    if clean_title in clean_first or clean_first in clean_title:
        return "\n".join(lines[1:]).strip()
    return text.strip()

# --- LLM CALLER ---
def call_llm(target_model: str, system_prompt: str, user_prompt: str, temp: float = 0.4, attempt: int = 1) -> str:
    current_model = target_model
    
    if attempt == 2:
        current_model = BACKUP_MODEL
        print(f"   >>> Grok failed. Switching to BACKUP: {current_model}")
    elif attempt > 2:
        return f"Error: Both AI models failed. Please try again later."

    try:
        api_key = os.environ.get("OPENROUTER_API_KEY")
        timeout = 60.0 
        
        system_prompt += " Do NOT use code blocks. Output raw Markdown only."

        with httpx.Client(timeout=timeout) as client:
            response = client.post(
                url="https://openrouter.ai/api/v1/chat/completions",
                headers={
                    "Authorization": f"Bearer {api_key}", 
                    "Content-Type": "application/json",
                    "HTTP-Referer": "http://localhost:5000",
                    "X-Title": "ScholarForge"
                },
                json={
                    "model": current_model,
                    "messages": [
                        {"role": "system", "content": system_prompt}, 
                        {"role": "user", "content": user_prompt}
                    ],
                    "temperature": temp,
                    "max_tokens": 4000
                }
            )
            
            if response.status_code != 200:
                print(f"   [!] AI Error ({current_model}): {response.status_code}")
                return call_llm(target_model, system_prompt, user_prompt, temp, attempt + 1)
                
            return clean_ai_output(response.json()['choices'][0]['message']['content'])
            
    except Exception as e:
        print(f"   [!] Exception ({current_model}): {e}")
        return call_llm(target_model, system_prompt, user_prompt, temp, attempt + 1)

# --- TASKS ---

def generate_summary(search_content: str, topic: str) -> str:
    return call_llm(
        SMART_MODEL,
        "You are a Senior Research Analyst.",
        f"Topic: {topic}\n\nRaw Data:\n{search_content[:15000]}\n\nTask: Summarize key facts, numbers, and trends."
    )

def generate_outline(topic: str, summary: str, format_type: str, target_pages: int) -> list:
    format_data = get_template_instructions(format_type, target_pages)
    prompt = (
        f"Topic: {topic}\nTarget: {format_data['target_sections']} sections.\n"
        f"Logic: {format_data['template_text']}\nContext: {summary[:2000]}\n"
        "Output: A JSON list of strings ONLY. Example: [\"1. Intro\", \"2. Body\"]"
    )
    content = call_llm(SMART_MODEL, "Return JSON only.", prompt, temp=0.2)
    
    match = re.search(r'\[.*\]', content.replace('\n', ' '), re.DOTALL)
    if match: return json.loads(match.group(0))
    return ["Introduction", "Analysis", "Conclusion"]

def generate_chart_from_data(summary: str, topic: str) -> str:
    try:
        chart_dir = "/app/static/charts"
        if not os.path.exists(chart_dir): os.makedirs(chart_dir, exist_ok=True)
        
        clean_name = re.sub(r'\W+', '', topic)[:15] 
        filename = f"chart_{clean_name}_{os.urandom(4).hex()}.png"
        filepath = os.path.join(chart_dir, filename)

        prompt = (
            f"Topic: {topic}\nContext: {summary[:3000]}\n"
            "Extract key trends/stats. ESTIMATE values if needed.\n"
            "Return JSON: {\"title\": \"...\", \"x_label\": \"...\", \"y_label\": \"...\", \"data\": [{\"label\": \"A\", \"value\": 10}]}"
        )
        
        content = call_llm(SMART_MODEL, "Return JSON only.", prompt, temp=0.1)
        match = re.search(r'\{.*\}', content.replace('\n', ' '), re.DOTALL)
        if not match: return None
        
        chart_data = json.loads(match.group(0))
        if not chart_data or 'data' not in chart_data: return None

        df = pd.DataFrame(chart_data['data'])
        
        # FIX: Use Object-Oriented Matplotlib interface for Thread Safety in Celery
        fig, ax = plt.subplots(figsize=(10, 6))
        plt.style.use('ggplot')
        
        ax.bar(df['label'], df['value'], color='#4f46e5', alpha=0.8)
        ax.set_title(chart_data.get('title', 'Analysis'), fontsize=14, pad=20)
        ax.set_xlabel(chart_data.get('x_label', ''), fontsize=12)
        ax.set_ylabel(chart_data.get('y_label', ''), fontsize=12)
        
        # Rotate x labels nicely
        plt.setp(ax.get_xticklabels(), rotation=45, ha='right')
        
        fig.tight_layout()
        fig.savefig(filepath, dpi=100)
        plt.close(fig) # Explicitly close figure to free memory
        
        return filepath
    except Exception as e:
        print(f"Chart Gen Error: {e}")
        return None

def critique_and_refine(section_text: str, topic: str) -> str:
    critic_prompt = (
        f"Topic: {topic}\nDraft:\n{section_text[:2000]}\n"
        "Identify ONE specific missing statistic. Return ONLY the search query. If good, return 'Pass'."
    )
    critique = call_llm(SMART_MODEL, "You are a harsh Editor.", critic_prompt, temp=0.1)
    
    if "Pass" in critique or "Error" in critique or len(critique) > 80: return section_text 
    
    new_data = get_search_results(critique, max_results=2)
    if "Error" in new_data or "No results" in new_data: return section_text

    refine_prompt = (
        f"Draft:\n{section_text}\n\nNew Verified Data:\n{new_data[:1500]}\n"
        "Integrate this new data naturally. Maintain Markdown."
    )
    return call_llm(SMART_MODEL, "You are a Senior Editor.", refine_prompt, temp=0.3)

def write_section(section_title: str, topic: str, summary: str, full_report_context: str, word_limit: int) -> str:
    base_prompt = f"Write a detailed report section '{section_title}' for a report on '{topic}'. Use research: {summary}. Length: {word_limit} words."
    
    keywords_for_table = ['comparison', 'market', 'financial', 'analysis', 'growth', 'impact', 'forecast', 'roi', 'cost']
    if any(k in section_title.lower() for k in keywords_for_table):
        base_prompt += "\n\nIMPORTANT: You MUST include a Markdown table comparing key metrics in this section."

    content = call_llm(SMART_MODEL, "You are a Report Writer. Use Markdown.", base_prompt, temp=0.5)
    
    if word_limit > 400 and "Error" not in content:
        content = critique_and_refine(content, topic)
        
    return clean_section_output(content, section_title)

# --- SCRAPING ---
def _get_article_text(url: str) -> str:
    try:
        headers = {'User-Agent': 'Mozilla/5.0'}
        with httpx.Client(timeout=15.0, follow_redirects=True) as client:
            response = client.get(url, headers=headers)
        if response.status_code != 200: return ""
        
        # Check for PDF
        if "application/pdf" in response.headers.get("Content-Type", "") or url.endswith(".pdf"):
            try:
                # fitz.open with stream requires "filetype" hint
                with fitz.open(stream=response.content, filetype="pdf") as doc:
                    text = ""
                    for i, page in enumerate(doc):
                        if i > 5: break 
                        text += page.get_text()
                return f"--- PDF SOURCE ---\n{text[:4000]}\n---"
            except: return ""
            
        soup = BeautifulSoup(response.text, 'lxml')
        for tag in soup(['script', 'style', 'nav', 'footer']): tag.decompose()
        return soup.get_text(separator='\n', strip=True)[:4000]
    except: return ""

def get_search_results(query: str, max_results: int = SEARCH_RESULTS_COUNT) -> str:
    try:
        api_key = os.environ.get("SERPAPI_KEY") 
        if not api_key: return "Error: SERPAPI_KEY not set."
        client = serpapi.Client(api_key=api_key)
        results = client.search({"q": query, "location": "US", "hl": "en", "gl": "us", "num": 5, "engine": "google"})
        snippets = []
        if "organic_results" in results:
            for i, result in enumerate(results["organic_results"]):
                url = result.get("link", "")
                title = result.get('title', '')
                snippet = result.get("snippet", "")
                full = ""
                if url and max_results > 3 and i < MAX_RESULTS_TO_SCRAPE: 
                    raw = _get_article_text(url)
                    if raw: full = f"\n[Full]: {raw[:1500]}"
                snippets.append(f"Source: {title}\nURL: {url}\nSummary: {snippet}{full}")
        return "\n\n".join(snippets) if snippets else "No results."
    except Exception as e: return f"Search Error: {e}"

# --- MAIN ORCHESTRATOR ---
def run_ai_engine_with_return(query: str, user_format: str, page_count: int = 15, task=None) -> tuple[str, str, str]: 
    def _update_status(message: str):
        print(message) 
        if task: task.update_state(state='PROGRESS', meta={'message': message})

    if not query: return "No query.", "", None

    _update_status("Step 1/6: Global Search (Deep Reading)...")
    search_content = get_search_results(query)
    
    _update_status("Step 2/6: Synthesizing...")
    summary = generate_summary(search_content, query)
    
    _update_status("Step 3/6: Visualizing Data...")
    chart_path = generate_chart_from_data(summary, query)
    
    _update_status("Step 4/6: Planning Structure...")
    outline = generate_outline(query, summary, user_format, page_count)

    total_words = page_count * WORDS_PER_PAGE 
    words_per_section = max(300, int(total_words / max(1, len(outline))))
    
    full_report = f"# {query.upper()}\n\n"
    for i, section in enumerate(outline):
        _update_status(f"Step 5/6: Researching & Writing {i+1}/{len(outline)}...")
        section_content = write_section(section, query, summary, full_report, words_per_section)
        full_report += f"\n\n## {section}\n{section_content}\n"
    
    _update_status("Step 6/6: Finalizing...")
    full_report = clean_ai_output(full_report)
    
    return search_content, full_report, chart_path

# --- CONVERTERS ---

def convert_to_txt(content, path):
    with open(path, "w", encoding="utf-8") as f: f.write(content)
    return "Success"

def convert_to_json(content, topic, path):
    data = {"topic": topic, "content": content, "generated_by": "ScholarForge"}
    with open(path, "w", encoding="utf-8") as f: json.dump(data, f, indent=4)
    return "Success"

def add_formatted_text(paragraph, text):
    """Splits text by ** and applies Bold/Black styling."""
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            clean_part = part[2:-2]
            run = paragraph.add_run(clean_part)
            run.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0) # Force Black
        else:
            paragraph.add_run(part)

def _add_markdown_table_to_docx(doc, table_block):
    lines = [l.strip() for l in table_block.strip().split('\n') if l.strip()]
    if len(lines) < 3: return
    
    rows = []
    for line in lines:
        if '---' in line: continue 
        clean_line = line.strip('|')
        cells = [c.strip() for c in clean_line.split('|')]
        # Clean asterisks from inside table cells
        cells = [c.replace('**', '') for c in cells]
        rows.append(cells)
    if not rows: return

    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = 'Table Grid'
    
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_data in enumerate(row_data):
            if c_idx < len(table.columns):
                cell = table.cell(r_idx, c_idx)
                cell.text = cell_data
                if r_idx == 0:
                    for run in cell.paragraphs[0].runs:
                        run.bold = True

def convert_to_docx(content, topic, path, chart_path=None):
    doc = Document()
    doc.add_heading(topic, 0)
    
    if chart_path and os.path.exists(chart_path):
        try: 
            doc.add_picture(chart_path, width=Inches(6))
            doc.add_paragraph("Figure 1: Analysis", style='Caption')
            doc.add_paragraph("") 
        except: pass

    lines = content.split('\n')
    table_buffer = []
    in_table = False
    
    for line in lines:
        stripped = line.strip()
        
        # Table Logic
        if stripped.startswith('|') and '|' in stripped[1:]:
            in_table = True
            table_buffer.append(stripped)
            continue
        elif in_table:
            _add_markdown_table_to_docx(doc, "\n".join(table_buffer))
            table_buffer = []
            in_table = False
        
        if not stripped: continue
        
        if stripped.startswith('## '):
            doc.add_heading(stripped.replace('## ', ''), level=2)
        elif stripped.startswith('# '):
            doc.add_heading(stripped.replace('# ', ''), level=1)
        elif stripped.startswith('### '):
            doc.add_heading(stripped.replace('### ', ''), level=3)
        elif stripped.startswith('#### '):
            doc.add_heading(stripped.replace('#### ', ''), level=4)
        elif stripped.startswith('* ') or stripped.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text(p, stripped[2:])
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6) 
            add_formatted_text(p, stripped)

    if in_table and table_buffer:
         _add_markdown_table_to_docx(doc, "\n".join(table_buffer))

    doc.save(path)
    return "Success"

def format_pdf_text(text):
    text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
    return re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', text)

def convert_to_pdf(content, topic, path, chart_path=None):
    doc = SimpleDocTemplate(path, pagesize=A4, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=18)
    styles = getSampleStyleSheet()
    story = []
    
    title = Paragraph(topic, styles['Title'])
    story.append(title)
    story.append(Spacer(1, 12))
    
    if chart_path and os.path.exists(chart_path):
        try: 
            story.append(RLImage(chart_path, width=450, height=250))
            story.append(Spacer(1, 12))
        except: pass

    normal_style = styles['Normal']
    normal_style.spaceAfter = 6
    
    lines = content.split('\n')
    table_buffer = []
    in_table = False

    for line in lines:
        stripped = line.strip()
        
        if stripped.startswith('|') and '|' in stripped[1:]:
            in_table = True
            table_buffer.append(stripped)
            continue
        elif in_table:
            data = []
            for row in table_buffer:
                if '---' in row: continue
                cells = [c.strip() for c in row.strip('|').split('|')]
                cells = [c.replace('**', '') for c in cells]
                data.append(cells)
            
            if data:
                t = Table(data)
                t.setStyle(TableStyle([
                    ('GRID', (0,0), (-1,-1), 1, colors.black),
                    ('BACKGROUND', (0,0), (-1,0), colors.lightgrey),
                    ('VALIGN', (0,0), (-1,-1), 'TOP'),
                ]))
                story.append(t)
                story.append(Spacer(1, 12))
            
            table_buffer = []
            in_table = False
        
        if not stripped: continue

        clean_text = format_pdf_text(stripped)

        if stripped.startswith('##'):
            story.append(Paragraph(clean_text.replace('#', ''), styles['Heading2']))
        elif stripped.startswith('####'):
            story.append(Paragraph(clean_text.replace('####', ''), styles['Heading3'])) 
        elif stripped.startswith('*') or stripped.startswith('-'):
            story.append(Paragraph(f"• {clean_text[2:]}", normal_style))
        else:
            story.append(Paragraph(clean_text, normal_style))
            
    doc.build(story)
    return "Success"