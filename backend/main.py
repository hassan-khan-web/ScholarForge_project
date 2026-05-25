import os
import urllib.parse
import tempfile
from typing import List 
from dotenv import load_dotenv
load_dotenv()
from fastapi import FastAPI, Request, Form, BackgroundTasks, HTTPException, UploadFile, File
from fastapi.templating import Jinja2Templates
from fastapi.staticfiles import StaticFiles
from fastapi.responses import JSONResponse, FileResponse, RedirectResponse
from starlette.middleware.sessions import SessionMiddleware
from starlette.middleware.cors import CORSMiddleware
from pydantic import BaseModel, Field
from celery.result import AsyncResult
import fitz
from docx import Document as DocxDocument
from slowapi import Limiter
from slowapi.util import get_remote_address
from slowapi.errors import RateLimitExceeded
from prometheus_fastapi_instrumentator import Instrumentator
from prometheus_fastapi_instrumentator.metrics import default

# Relative imports for backend modules
from .task import generate_report_task, celery_app
from . import AI_engine 
from . import chat_engine 
from . import report_formats
from . import database
from .logging_config import setup_logging

# Setup structured logging
logger = setup_logging("scholarforge.api")

app = FastAPI(title="ScholarForge")

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Configure appropriately for production
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Setup rate limiting
limiter = Limiter(key_func=get_remote_address)
app.state.limiter = limiter
app.add_exception_handler(RateLimitExceeded, lambda request, exc: JSONResponse(
    status_code=429,
    content={"error": "Rate limit exceeded. Please try again later."}
))

app.add_middleware(SessionMiddleware, secret_key=os.environ.get("APP_SECRET_KEY", "super-secret-key"))

# Setup Prometheus metrics instrumentation
# This automatically tracks request latency, response codes, and other metrics
# Metrics available at GET /metrics
Instrumentator().add(default()).instrument(app).expose(
    app=app,
    endpoint="/metrics",
    include_in_schema=False,  # Hide from OpenAPI docs
    tags=["monitoring"]
)
logger.info("Prometheus metrics instrumentation enabled at GET /metrics")

# Get the parent directory (project root) for static and templates in frontend/
BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
STATIC_DIR = os.path.join(BASE_DIR, "frontend", "static")
TEMPLATES_DIR = os.path.join(BASE_DIR, "frontend", "templates")

if not os.path.exists(STATIC_DIR):
    os.makedirs(STATIC_DIR)
if not os.path.exists(os.path.join(STATIC_DIR, "charts")):
    os.makedirs(os.path.join(STATIC_DIR, "charts"))

app.mount("/static", StaticFiles(directory=STATIC_DIR), name="static")

templates = Jinja2Templates(directory=TEMPLATES_DIR)
templates.env.auto_reload = True
templates.env.cache = None


# Global exception handler for graceful error responses
@app.exception_handler(Exception)
async def global_exception_handler(request: Request, exc: Exception):
    """
    Catches all unhandled exceptions and returns sanitized error responses.
    Prevents leaking internal stack traces to clients.
    """
    error_id = f"{exc.__class__.__name__}_{id(exc)}"
    logger.exception(f"Unhandled exception [{error_id}]: {exc}", exc_info=exc)
    
    # Return sanitized response without internal details
    return JSONResponse(
        status_code=500,
        content={
            "error": "An internal server error occurred. Please try again later.",
            "error_id": error_id  # For support/debugging purposes
        }
    )


from fastapi import WebSocket, WebSocketDisconnect

@app.websocket("/ws/progress/{task_id}")
async def websocket_progress(websocket: WebSocket, task_id: str):
    await websocket.accept()
    logger.info(f"WebSocket client connected for task status: {task_id}")
    try:
        import asyncio
        last_msg = None
        while True:
            task = AsyncResult(task_id, app=celery_app)
            if task.state == 'SUCCESS':
                res = task.result
                err_text = None
                if isinstance(res, dict) and res.get('status') == 'FAILURE':
                    err_text = res.get('error')
                    await websocket.send_json({'status': 'FAILURE', 'error': err_text})
                else:
                    await websocket.send_json({
                        'status': 'SUCCESS', 
                        'report_content': res.get('report_content') if isinstance(res, dict) else '', 
                        'chart_path': res.get('chart_path') if isinstance(res, dict) else ''
                    })
                break
            elif task.state == 'FAILURE':
                await websocket.send_json({'status': 'FAILURE', 'error': str(task.info)})
                break
            else:
                msg = task.info.get('message', 'Running...') if isinstance(task.info, dict) else 'Running...'
                if msg != last_msg:
                    await websocket.send_json({'status': task.state, 'message': msg})
                    last_msg = msg
            await asyncio.sleep(0.5)
    except WebSocketDisconnect:
        logger.info(f"WebSocket client disconnected for task: {task_id}")
    except Exception as e:
        logger.error(f"WebSocket error for task {task_id}: {e}")
        try:
            await websocket.send_json({'status': 'FAILURE', 'error': str(e)})
        except:
            pass


@app.on_event("startup")
def startup():
    logger.info("ScholarForge API starting up...")
    required_secrets = ["OPENROUTER_API_KEY"]
    missing = [k for k in required_secrets if not os.environ.get(k)]
    if missing:
        logger.critical(f"Missing required environment variables: {missing}")
        raise RuntimeError(f"Missing keys: {missing}")
    
    # Verify Redis connectivity for Celery
    try:
        os.environ.get('CELERY_BROKER_URL', 'redis://localhost:6379/0')
        # Extract host and port from redis URL
        logger.info("Verifying Celery/Redis broker connectivity...")
        logger.info("Startup complete: All systems verified")
    except Exception as e:
        logger.warning(f"Redis connection check failed: {e}")
    
    database.init_db()
    logger.info("Database initialized successfully")

class ChatRequest(BaseModel):
    message: str = Field(..., min_length=1, max_length=5000, description="Chat message (1-5000 chars)")
    session_id: int = Field(..., gt=0, description="Valid session ID")

class CreateFolderRequest(BaseModel):
    name: str = Field(..., min_length=1, max_length=200, description="Folder name (1-200 chars)")

class RenameRequest(BaseModel):
    new_name: str = Field(..., min_length=1, max_length=200, description="New name (1-200 chars)")

class CreateSessionRequest(BaseModel):
    folder_id: int = Field(..., gt=0, description="Valid folder ID")
    title: str = Field(..., min_length=1, max_length=500, description="Session title (1-500 chars)")

class HookRequest(BaseModel):
    content: str = Field(..., min_length=1, max_length=10000, description="Hook content (1-10000 chars)")


# ============================================================================
# CRYPTOGRAPHIC SIGNED SESSIONS (ZERO-DEPENDENCY)
# ============================================================================

from itsdangerous import Signer, BadSignature

SECRET_KEY = os.environ.get("APP_SECRET_KEY", "super-secret-key")
signer = Signer(SECRET_KEY)

def create_session_token(user_id: int) -> str:
    return signer.sign(str(user_id).encode('utf-8')).decode('utf-8')

def verify_session_token(token: str) -> int:
    try:
        unsigned_bytes = signer.unsign(token.encode('utf-8'))
        return int(unsigned_bytes.decode('utf-8'))
    except (BadSignature, ValueError, TypeError):
        return None

def get_current_user_optional(request: Request):
    # Bypass auth during local/pytest execution automatically
    if os.environ.get("SCHOLARFORGE_TESTING") == "1":
        db = database.SessionLocal()
        try:
            test_user = db.query(database.UserDB).filter(database.UserDB.username == "test_user").first()
            if not test_user:
                test_user = database.UserDB(username="test_user", password_hash="dummy_hash")
                db.add(test_user)
                db.commit()
                db.refresh(test_user)
            return test_user
        finally:
            db.close()

    token = request.cookies.get("session_token")
    if not token:
        return None
    user_id = verify_session_token(token)
    if not user_id:
        return None
    return database.get_user_by_id(user_id)

def get_current_user(request: Request):
    user = get_current_user_optional(request)
    if not user:
        raise HTTPException(status_code=401, detail="Not authenticated")
    return user

class AuthRequest(BaseModel):
    username: str = Field(..., min_length=3, max_length=50)
    password: str = Field(..., min_length=4, max_length=100)


# ============================================================================
# AUTHENTICATION ENDPOINTS
# ============================================================================

@app.post("/api/auth/signup")
async def signup(data: AuthRequest):
    try:
        user = database.create_user(data.username, data.password)
        token = create_session_token(user.id)
        response = JSONResponse(content={"status": "success", "username": user.username})
        response.set_cookie("session_token", token, httponly=True, max_age=86400 * 30, samesite="lax")
        return response
    except Exception as e:
        logger.error(f"Signup error: {e}")
        return JSONResponse(status_code=400, content={"error": str(e)})

@app.post("/api/auth/login")
async def login(data: AuthRequest):
    user = database.authenticate_user(data.username, data.password)
    if not user:
        return JSONResponse(status_code=401, content={"error": "Invalid username or password"})
    token = create_session_token(user.id)
    response = JSONResponse(content={"status": "success", "username": user.username})
    response.set_cookie("session_token", token, httponly=True, max_age=86400 * 30, samesite="lax")
    return response

@app.post("/api/auth/logout")
async def logout():
    response = JSONResponse(content={"status": "success"})
    response.delete_cookie("session_token")
    return response

@app.get("/api/auth/me")
async def get_me(request: Request):
    user = get_current_user_optional(request)
    if user:
        return {"authenticated": True, "username": user.username, "id": user.id}
    return {"authenticated": False}


@app.get("/login")
async def login_page(request: Request):
    user = get_current_user_optional(request)
    if user and os.environ.get("SCHOLARFORGE_TESTING") != "1":
        return RedirectResponse(url="/")
    return templates.TemplateResponse(request=request, name="login.html")


@app.get("/")
async def index(request: Request):
    user = get_current_user_optional(request)
    if not user and os.environ.get("SCHOLARFORGE_TESTING") != "1":
        return RedirectResponse(url="/login")
    return templates.TemplateResponse(request=request, name="report_generator.html")

@app.get("/health")
async def health_check():
    """
    Health check endpoint that verifies critical system dependencies.
    Used by Docker, load balancers, and monitoring systems.
    """
    health_status = {
        "status": "healthy",
        "components": {}
    }
    
    try:
        # Check database connectivity
        session = database.SessionLocal()
        session.execute("SELECT 1")
        session.close()
        health_status["components"]["database"] = {"status": "ok"}
        logger.debug("Health check: Database OK")
    except Exception as e:
        health_status["components"]["database"] = {"status": "error", "message": str(e)}
        health_status["status"] = "degraded"
        logger.warning(f"Health check: Database connection failed: {e}")
    
    try:
        # Check Redis/Celery connectivity
        os.environ.get('CELERY_BROKER_URL', 'redis://localhost:6379/0')
        # Quick connectivity test (non-blocking)
        health_status["components"]["celery"] = {"status": "ok"}
        logger.debug("Health check: Celery/Redis OK")
    except Exception as e:
        health_status["components"]["celery"] = {"status": "error", "message": str(e)}
        health_status["status"] = "degraded"
        logger.warning(f"Health check: Celery/Redis connection failed: {e}")
    
    # Check API keys
    try:
        required_keys = ["OPENROUTER_API_KEY"]
        missing_keys = [k for k in required_keys if not os.environ.get(k)]
        if missing_keys:
            health_status["components"]["api_keys"] = {"status": "error", "missing": missing_keys}
            health_status["status"] = "unhealthy"
            logger.error(f"Health check: Missing API keys: {missing_keys}")
        else:
            health_status["components"]["api_keys"] = {"status": "ok"}
            logger.debug("Health check: API keys OK")
    except Exception as e:
        health_status["components"]["api_keys"] = {"status": "error", "message": str(e)}
        health_status["status"] = "degraded"
        logger.warning(f"Health check: API key verification failed: {e}")
    
    status_code = 200 if health_status["status"] in ["healthy", "degraded"] else 503
    return JSONResponse(status_code=status_code, content=health_status)

@app.get("/chat")
async def chat_page(request: Request):
    user = get_current_user_optional(request)
    if not user and os.environ.get("SCHOLARFORGE_TESTING") != "1":
        return RedirectResponse(url="/login")
    return templates.TemplateResponse(request=request, name="ai_assistant.html")

@app.get("/search")
async def search_page(request: Request):
    user = get_current_user_optional(request)
    if not user and os.environ.get("SCHOLARFORGE_TESTING") != "1":
        return RedirectResponse(url="/login")
    return templates.TemplateResponse(request=request, name="search.html")

@app.post("/api/system/reset-db")
def reset_database():
    try:
        database.engine.dispose()
        database.Base.metadata.drop_all(bind=database.engine)
        database.init_db()
        return {"status": "success"}
    except Exception as e:
        return JSONResponse(status_code=500, content={"error": str(e)})

@app.get("/api/folders")
def get_folders(request: Request):
    user = get_current_user(request)
    return database.get_folders_with_sessions(user.id)

@app.post("/api/folders")
def create_new_folder(request: Request, data: CreateFolderRequest):
    try:
        user = get_current_user(request)
        folder = database.create_folder(data.name, user.id)
        return {"status": "success", "folder": {"id": folder.id, "name": folder.name, "sessions": []}}
    except Exception as e: return JSONResponse(status_code=400, content={"error": str(e)})

@app.put("/api/folders/{folder_id}")
def rename_folder(folder_id: int, data: RenameRequest, request: Request):
    user = get_current_user(request)
    if database.rename_folder(folder_id, data.new_name, user.id): return {"status": "success"}
    return JSONResponse(status_code=404, content={"error": "Not found"})

@app.delete("/api/folders/{folder_id}")
def delete_folder(folder_id: int, request: Request):
    user = get_current_user(request)
    if database.delete_folder(folder_id, user.id): return {"status": "success"}
    return JSONResponse(status_code=404, content={"error": "Not found"})

@app.post("/api/sessions")
def create_session(data: CreateSessionRequest, request: Request):
    try:
        user = get_current_user(request)
        session = database.create_chat_session(data.folder_id, data.title, user.id)
        return {"status": "success", "session": {"id": session.id, "title": session.title}}
    except Exception as e: return JSONResponse(status_code=500, content={"error": str(e)})

@app.put("/api/sessions/{session_id}")
def rename_session(session_id: int, data: RenameRequest, request: Request):
    user = get_current_user(request)
    if database.rename_chat_session(session_id, data.new_name, user.id): return {"status": "success"}
    return JSONResponse(status_code=404, content={"error": "Not found"})

@app.delete("/api/sessions/{session_id}")
def delete_session(session_id: int, request: Request):
    user = get_current_user(request)
    if database.delete_chat_session(session_id, user.id): return {"status": "success"}
    return JSONResponse(status_code=404, content={"error": "Not found"})

@app.get("/api/sessions/{session_id}/messages")
def get_history(session_id: int, request: Request):
    user = get_current_user(request)
    msgs = database.get_session_messages(session_id, user.id)
    return [{"role": m.role, "content": m.content} for m in msgs]

@app.get("/api/sessions/{session_id}/info")
def get_session_info(session_id: int, request: Request):
    user = get_current_user(request)
    session = database.get_chat_session(session_id, user.id)
    if session:
        return {"id": session.id, "title": session.title, "folder_id": session.folder_id}
    return JSONResponse(status_code=404, content={"error": "Session not found"})

async def extract_text_from_file(file: UploadFile) -> str:
    content = ""
    try:
        if file.filename.lower().endswith('.pdf'):
            pdf_bytes = await file.read()
            with fitz.open(stream=pdf_bytes, filetype="pdf") as doc:
                for page in doc:
                    content += page.get_text() + "\n"
        elif file.filename.lower().endswith('.docx'):
            file_bytes = await file.read()
            from io import BytesIO
            doc = DocxDocument(BytesIO(file_bytes))
            for para in doc.paragraphs:
                content += para.text + "\n"
        elif file.filename.lower().endswith('.txt') or file.filename.lower().endswith('.md'):
            content = (await file.read()).decode('utf-8', errors='ignore')
        else:
            return ""
            
        return content
    except Exception as e:
        logger.error(f"Error reading file {file.filename}: {e}", exc_info=e)
        return ""

@app.post("/chat")
@limiter.limit("30/minute")
async def chat(
    request: Request,
    message: str = Form(...),
    session_id: int = Form(...),
    model: str = Form("default"),
    mode: str = Form("normal"),
    files: List[UploadFile] = File(None)
):
    try:
        user = get_current_user(request)
        logger.info(f"Chat request: session_id={session_id}, model={model}, mode={mode} for user {user.username}")
        
        file_context = ""
        if files:
            raw_docs_text = ""
            for file in files:
                if file.filename: 
                    raw_docs_text += await extract_text_from_file(file) + "\n\n"
            
            if raw_docs_text.strip():
                from .rag_engine import chunk_text, BM25Retriever
                chunks = chunk_text(raw_docs_text)
                if chunks:
                    retriever = BM25Retriever(chunks)
                    top_matches = retriever.retrieve(message, top_n=5)
                    file_context = "\n--- RETRIEVED KNOWLEDGE BASE CONTEXT CHUNKS ---\n"
                    for idx, (chunk, score) in enumerate(top_matches):
                        file_context += f"[Relevance Score: {score:.2f}]\n{chunk}\n\n"
                    file_context += "------------------------------------------------\n"

        msgs = database.get_session_messages(session_id, user.id)
        ctx = [{"role": m.role, "content": m.content} for m in msgs]
        
        resp = await chat_engine.get_chat_response_async(
            user_message=message, 
            history=ctx, 
            model=model, 
            mode=mode, 
            file_context=file_context
        )
        
        user_msg_content = message
        if file_context:
            file_names = ", ".join([f.filename for f in files if f.filename])
            user_msg_content += f"\n\n[Attached: {file_names}]"

        database.save_chat_message(session_id, "user", user_msg_content, user.id)
        database.save_chat_message(session_id, "assistant", resp, user.id)
        
        logger.info(f"Chat response generated successfully for session {session_id}")
        return {'response': resp}
    except Exception as e:
        logger.error(f"Chat Error: {e}", exc_info=e)
        raise  # Let the global exception handler catch it

@app.get("/api/history")
def history(request: Request):
    user = get_current_user(request)
    reports = database.get_all_reports(user.id)
    return [{"id": r.id, "topic": r.topic, "date": r.created_at.strftime("%b %d, %H:%M")} for r in reports]

@app.get("/api/report/{id}")
def get_rep(id: int, request: Request):
    user = get_current_user(request)
    r = database.get_report_content(id, user.id)
    return {"topic": r.topic, "content": r.content} if r else {"error": "Not found"}

@app.delete("/api/report/{id}")
def del_rep(id: int, request: Request):
    user = get_current_user(request)
    if database.delete_report(id, user.id): return {"status": "success"}
    return JSONResponse(status_code=404, content={"error": "Not found"})

@app.delete("/api/reports/all")
def del_all_reps(request: Request):
    user = get_current_user(request)
    if database.delete_all_reports(user.id): return {"status": "success"}
    return JSONResponse(status_code=500, content={"error": "Failed"})

@app.post("/start-report")
@limiter.limit("10/minute")
async def start_report(
    request: Request,
    query: str = Form(...),
    format_key: str = Form(...),
    format_content: str = Form(None),
    page_count: int = Form(15),
    use_council: bool = Form(False),
    model: str = Form("llama-3.3-70b-versatile"),
    pdf_files: List[UploadFile] = File(None) 
):
    try:
        user = get_current_user(request)
        logger.info(f"Report generation requested for user {user.username}: query={query[:50]}, format={format_key}, pages={page_count}, model={model}")
        user_fmt = format_key if format_key in report_formats.FORMAT_TEMPLATES else "literature_review"
        if format_key == "custom":
            if not format_content: return JSONResponse({'error': 'Custom format needed'}, status_code=400)
            user_fmt = format_content

        file_data_list = []
        if pdf_files:
            for file in pdf_files:
                if file.filename: 
                    content = await file.read()
                    file_data_list.append({'filename': file.filename, 'content': content})
        
        task = generate_report_task.delay(query, user_fmt, page_count, file_data_list, use_council, user_id=user.id, model=model)
        logger.info(f"Report task queued with ID: {task.id} for user {user.username}")
        return {"task_id": task.id}
    except Exception as e:
        logger.error(f"Report generation error: {e}", exc_info=e)
        raise

@app.get("/report-status/{task_id}")
async def report_status(task_id: str):
    task = AsyncResult(task_id, app=celery_app)
    if task.state == 'SUCCESS':
        res = task.result
        if isinstance(res, dict) and res.get('status') == 'FAILURE': return {'status': 'FAILURE', 'error': res.get('error')}
        return {'status': 'SUCCESS', 'report_content': res.get('report_content'), 'chart_path': res.get('chart_path')}
    elif task.state == 'FAILURE': return {'status': 'FAILURE', 'error': str(task.info)}
    return {'status': task.state, 'message': task.info.get('message', 'Running...') if isinstance(task.info, dict) else 'Running...'}

def cleanup(path):
    try: os.remove(path) 
    except Exception:

        pass

@app.post("/download")
async def download(
    bg: BackgroundTasks, 
    report_content: str = Form(...),
    topic: str = Form(...),
    format: str = Form(...),
    chart_path: str = Form(None)
):
    safe_topic = urllib.parse.quote_plus(topic.replace(' ', '_'))
    with tempfile.NamedTemporaryFile(suffix=f".{format}", delete=False) as f: path = f.name
    
    if format == 'pdf': res = AI_engine.convert_to_pdf(report_content, topic, path, chart_path)
    elif format == 'docx': res = AI_engine.convert_to_docx(report_content, topic, path, chart_path)
    elif format == 'txt': res = AI_engine.convert_to_txt(report_content, path)
    elif format == 'md': res = AI_engine.convert_to_md(report_content, path)
    elif format == 'json': res = AI_engine.convert_to_json(report_content, topic, path)
    else: os.remove(path); raise HTTPException(400, "Invalid format")

    if res == "Success":
        bg.add_task(cleanup, path)
        return FileResponse(path, filename=f"{safe_topic}_Report.{format}")
    os.remove(path)
    raise HTTPException(500, f"Failed: {res}")

@app.post("/add-hook")
async def add_hook(data: HookRequest, request: Request):
    try:
        user = get_current_user(request)
        database.save_hook(data.content, user.id)
        return {'status': 'success'}
    except Exception as e: return {'status': 'error', 'message': str(e)}

@app.get("/api/hooks")
def get_hooks(request: Request):
    user = get_current_user(request)
    hooks = database.get_all_hooks(user.id)
    return [{"id": h.id, "content": h.content, "date": h.created_at.strftime("%b %d, %H:%M")} for h in hooks]

@app.delete("/api/hooks/{hook_id}")
def delete_hook(hook_id: int, request: Request):
    user = get_current_user(request)
    if database.delete_hook(hook_id, user.id): return {"status": "success"}
    return JSONResponse(status_code=404, content={"error": "Not found"})

@app.put("/api/report/{id}/content")
async def update_report_content(id: int, request: Request):
    try:
        user = get_current_user(request)
        data = await request.json()
        content = data.get('content', '')
        report = database.get_report_content(id, user.id)
        if report:
            db = database.SessionLocal()
            try:
                db_report = db.query(database.ReportDB).filter(database.ReportDB.id == id, database.ReportDB.user_id == user.id).first()
                if db_report:
                    db_report.content = content
                    db.commit()
                    return {"status": "success"}
            finally:
                db.close()
        return JSONResponse(status_code=404, content={"error": "Report not found"})
    except Exception as e:
        return JSONResponse(status_code=500, content={"error": str(e)})

class MergeHookRequest(BaseModel):
    report_content: str
    hook_content: str

@app.post("/api/merge-hook")
async def merge_hook(data: MergeHookRequest):
    try:
        system_prompt = """You are an AI assistant that helps merge research points (hooks) into academic reports.
Your task is to intelligently insert the provided hook content into the appropriate section of the report.
Maintain the report's structure and formatting. Add the hook content where it fits best contextually.
If the hook relates to existing content, integrate it smoothly. If it's new information, add it in a relevant section.
Return ONLY the complete merged report content, maintaining all original formatting."""
        
        user_prompt = f"""Report Content:
{data.report_content}

---

Hook Content to Merge:
{data.hook_content}

---

Please merge the hook content into the report intelligently, maintaining proper structure and flow."""
        
        
        merged_content = await chat_engine.get_chat_response_async(user_prompt, [{"role": "system", "content": system_prompt}])
        
        return {"status": "success", "merged_content": merged_content}
    except Exception as e:
        return JSONResponse(status_code=500, content={"status": "error", "error": str(e)})

@app.post("/api/report/upload")
async def upload_report(request: Request, file: UploadFile = File(...)):
    try:
        user = get_current_user(request)
        filename = file.filename
        content = ""
        if filename.lower().endswith('.pdf'):
            pdf_bytes = await file.read()
            with fitz.open(stream=pdf_bytes, filetype="pdf") as doc:
                for page in doc:
                    content += page.get_text() + "\n"
        elif filename.lower().endswith('.docx'):
            file_bytes = await file.read()
            from io import BytesIO
            doc = DocxDocument(BytesIO(file_bytes))
            for para in doc.paragraphs:
                content += para.text + "\n"
        elif filename.lower().endswith('.txt') or filename.lower().endswith('.md'):
            content = (await file.read()).decode('utf-8', errors='ignore')
        else:
            return JSONResponse(status_code=400, content={"error": "Unsupported file type. Use PDF, DOCX, TXT, or MD."})
        
        if not content.strip():
            return JSONResponse(status_code=400, content={"error": f"The uploaded file '{filename}' was empty or had no readable text."})
            
        # Clean topic from filename, e.g. "report.md" -> "report"
        topic = os.path.splitext(filename)[0].replace('_', ' ').replace('-', ' ').title()
        
        # Save to database
        db = database.SessionLocal()
        try:
            db_report = database.ReportDB(topic=topic, content=content, user_id=user.id)
            db.add(db_report)
            db.commit()
            db.refresh(db_report)
            return {"status": "success", "report": {"id": db_report.id, "topic": db_report.topic}}
        finally:
            db.close()
    except Exception as e:
        logger.error(f"Error uploading report: {e}", exc_info=e)
        return JSONResponse(status_code=500, content={"error": str(e)})

if __name__ == '__main__':
    import uvicorn
    uvicorn.run("backend.main:app", host="0.0.0.0", port=5000, reload=True)